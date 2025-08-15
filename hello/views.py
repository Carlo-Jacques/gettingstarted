import os
import json
from docx import Document # pyright: ignore[reportMissingImports]
from datetime import datetime  # Added for timestamp
from django.conf import settings
from django.shortcuts import render
from django.http import JsonResponse
from django.core.mail import EmailMessage
from django.views.decorators.csrf import csrf_exempt

import os
import json
import traceback
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.core.mail import EmailMessage
from django.conf import settings
from django.shortcuts import render
from docx import Document # pyright: ignore[reportMissingImports]
from datetime import datetime

@csrf_exempt
def generate_docs(request):
    # Early logging to confirm view is reached
    debug_dir = os.path.join(settings.BASE_DIR, 'generated_docs')
    os.makedirs(debug_dir, exist_ok=True)
    early_log_path = os.path.join(debug_dir, 'early_debug_log.txt')
    try:
        with open(early_log_path, 'a') as log:
            log.write(f"== generate_docs called at {datetime.now()} ==\n")
    except Exception as e:
        return JsonResponse({'error': f'Failed to write early debug log: {str(e)}'}, status=500)

    if request.method != 'POST':
        return JsonResponse({'error': 'Invalid method'}, status=400)

    try:
        # Parse JSON payload
        try:
            data = json.loads(request.body.decode('utf-8'))
        except json.JSONDecodeError as e:
            return JsonResponse({'error': f'Invalid JSON payload: {str(e)}'}, status=400)

        if not data.get('email'):
            return JsonResponse({'error': 'Email address is required'}, status=400)

        generated_files = []
        output_dir = os.path.join(settings.BASE_DIR, 'generated_docs')
        os.makedirs(output_dir, exist_ok=True)

        # Unique subdirectory per request
        request_id = str(os.urandom(8).hex())
        user_output_dir = os.path.join(output_dir, request_id)
        os.makedirs(user_output_dir, exist_ok=True)

        # Log payload for debugging
        debug_log_path = os.path.join(user_output_dir, 'debug_log.txt')
        with open(debug_log_path, 'a') as log:
            log.write(f"== Django view started (Request ID: {request_id}, {datetime.now()}) ==\n")
            log.write(f"Received payload: {json.dumps(data, indent=2)}\n\n")
            log.write(f"BASE_DIR: {settings.BASE_DIR}\n")
            log.write(f"Template dir: {os.path.join(settings.BASE_DIR, 'hello', 'template_docs')}\n")

        # Generate timestamp for filenames
        timestamp = datetime.now().strftime('%Y%m%d%H%M%S')

        # Generate docx function
        def generate_docx(template, output_name, mapping, output_list):
            template_path = os.path.join(settings.BASE_DIR, 'hello', 'template_docs', template)
            if not os.path.exists(template_path):
                with open(debug_log_path, 'a') as log:
                    log.write(f"Template not found: {template_path}\n")
                return False
            try:
                doc = Document(template_path)
                def replace_text_preserving_format(paragraphs, mapping):
                    for para in paragraphs:
                        full_text = ''.join(run.text for run in para.runs)
                        replaced_text = full_text
                        for key, val in mapping.items():
                            replaced_text = replaced_text.replace(key, str(val))
                        if replaced_text != full_text:
                            for run in para.runs:
                                run.text = ''
                            if para.runs:
                                para.runs[0].text = replaced_text
                replace_text_preserving_format(doc.paragraphs, mapping)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            replace_text_preserving_format(cell.paragraphs, mapping)
                output_path = os.path.join(user_output_dir, output_name)
                doc.save(output_path)
                output_list.append(output_path)
                return True
            except Exception as e:
                with open(debug_log_path, 'a') as log:
                    log.write(f"Error generating {output_name}: {str(e)}\n{traceback.format_exc()}\n")
                return False

        # Normalize document_types
        doc_type = data.get("document_types", [])
        document_types = [doc_type.lower()] if isinstance(doc_type, str) else [d.lower() for d in doc_type]

        # Extract key fields
        first_name = data.get("first_name", "").upper()
        last_name = data.get("last_name", "").upper()
        full_name = f"{first_name} {last_name}"
        city = data.get("city", "")
        state = data.get("state", "")
        county = data.get("county", "")
        ZIP = data.get("zip", "")
        PRINCIPAL_ADDRESS_LINE_2 = f"{city}, {state} {ZIP}"

        replacements = {
            "{{FIRST_NAME}}": first_name,
            "{{LAST_NAME}}": last_name,
            "{{Principal_Full_Name}}": full_name,
            "{{Principal_City}}": city,
            "{{Principal_State}}": state,
            "{{Principal_Zip}}": ZIP,
            "{{Principal_Address_Line_2}}": PRINCIPAL_ADDRESS_LINE_2
        }

        # Will
        if "will" in document_types:
            will_replacements = replacements.copy()
            will_replacements.update({
                "{{TESTATORSFULLNAME}}": full_name,
                "{{TESTATORSCOUNTY}}": county,
                "{{COUNTY}}": county,
                "{{TESTATORSSTATE}}": state,
                "{{FATHER_MOTHER}}": data.get("father_mother", ""),
                "{{CHILDREN}}": data.get("CHILDREN", ""),
                "{{TESTATOR_TESTATRIX}}": data.get("testator_testatrix", ""),
                "{{TESTATORSADDRESSLINE1}}": data.get("address_line_1", ""),
                "{{TESTATORSADDRESSLINE2}}": PRINCIPAL_ADDRESS_LINE_2,
                "{{TESTATORSPHONENUMBER}}": data.get("phone", ""),
                "{{AGENTSFULLNAME1}}": (data.get("will_agent_1_full_name") or data.get("shared_agent_1_full_name") or "").upper(),
                "{{AGENTSADDRESS1}}": data.get("will_agent_1_address_1") or data.get("shared_agent_1_address_1", ""),
                "{{AGENTSCITY1}}": data.get("will_agent_1_city") or data.get("shared_agent_1_city", ""),
                "{{AGENTSCOUNTY1}}": data.get("will_agent_1_county") or data.get("shared_agent_1_county", ""),
                "{{AGENTSSTATE1}}": data.get("will_agent_1_state") or data.get("shared_agent_1_state", ""),
                "{{AGENTSPHONENUMBER1}}": data.get("will_agent_1_phone") or data.get("shared_agent_1_phone", ""),
                "{{AGENTSFULLNAME2}}": data.get("will_agent_2_full_name", "").upper(),
                "{{AGENTSADDRESS2}}": data.get("will_agent_2_address_1", ""),
                "{{AGENTSCITY2}}": data.get("will_agent_2_city", ""),
                "{{AGENTSCOUNTY2}}": data.get("will_agent_2_county", ""),
                "{{AGENTSSTATE2}}": data.get("will_agent_2_state", ""),
                "{{AGENTSPHONENUMBER2}}": data.get("will_agent_2_phone", ""),
                "{{Signing_Date}}": data.get("date", ""),
                "{{DATE}}": data.get("date", "")
            })
            if generate_docx(
                template="Template_Last Will & Testament.docx",
                output_name=f"WILL_{first_name}_{last_name}_{timestamp}.docx",
                mapping=will_replacements,
                output_list=generated_files
            ):
                with open(debug_log_path, 'a') as log:
                    log.write("Generated Will doc\n")
            else:
                with open(debug_log_path, 'a') as log:
                    log.write("Failed to generate Will doc\n")

        # Living Will
        def format_living_choice(initials_key, check_key, label):
            initials = data.get(initials_key, "").strip().upper()
            checked = data.get(check_key, "") in ["on", "true", True]
            check_symbol = "☑" if checked else "☐"
            return f"__{initials}__ {check_symbol} - {label}"

        def format_food_water_choice(initial_yes_key, initial_no_key, choice_key):
            initials_yes = data.get(initial_yes_key, "").strip().upper()
            initials_no = data.get(initial_no_key, "").strip().upper()
            choice = data.get(choice_key, "").lower()
            line_yes = f"__{initials_yes}__ {'☑' if choice == 'yes' else '☐'} - Even if I have the quality of life described above, I still wish to be treated with food and water by tube or intravenously (IV)."
            line_no = f"__{initials_no}__ {'☑' if choice == 'no' else '☐'} - If I have the quality of life described above, I do NOT wish to be treated with food and water by tube or intravenously (IV)."
            return f"{line_yes}\n{line_no}"

        def format_life_sustaining_choice(initial_key, check_key, label, extra_text_key=None):
            initials = data.get(initial_key, "").strip().upper()
            checked = data.get(check_key, "") in ["on", "true", True]
            check_symbol = "☑" if checked else "☐"
            extra = f": {data.get(extra_text_key, '').strip()}" if extra_text_key and data.get(extra_text_key) else ""
            return f"__{initials}__ {check_symbol} - {label}{extra}"

        if "living" in document_types:
            living_unacceptable = "\n".join([
                format_living_choice("lw_ls_initial_coma", "lw_ls_check_coma", "Chronic coma or persistent vegetative state"),
                format_living_choice("lw_ls_initial_nocomm", "lw_ls_check_nocomm", "No longer able to communicate my needs"),
                format_living_choice("lw_ls_initial_recff", "lw_ls_check_recff", "No longer able to recognize family or friends"),
                format_living_choice("lw_ls_initial_totdep", "lw_check_totdep", "Total dependence on others for daily care"),
                format_living_choice("lw_ls_initial_other", "lw_ls_check_other", f"Other: {data.get('lw_ls_text_other', '')}")
            ])
            living_life_sustaining = "\n".join([
                format_life_sustaining_choice("lw_cpr_initial", "lw_cpr", "Cardiopulmonary Resuscitation (CPR)"),
                format_life_sustaining_choice("lw_vent_initial", "lw_vent", "Ventilation (breathing machine)"),
                format_life_sustaining_choice("lw_feed_initial", "lw_feed", "Feeding tube"),
                format_life_sustaining_choice("lw_dialysis_initial", "lw_dialysis", "Dialysis"),
                format_life_sustaining_choice("lw_other_initial", "lw_other", "Other", "lw_other_text")
            ])
            living_food_water = format_food_water_choice(
                "lw_foodwater_initial_yes",
                "lw_foodwater_initial_no",
                "lw_foodwater_choice"
            )
            living_replacements = replacements.copy()
            living_replacements.update({
                "{{LW_UNACCEPTABLE_LIFE}}": living_unacceptable,
                "{{LW_FOOD_WATER_IV}}": living_food_water,
                "{{LW_LIFE_SUSTAINING}}": living_life_sustaining
            })
            if generate_docx(
                template="Template_Living_Will.docx",
                output_name=f"Living_Will_{first_name}_{last_name}_{timestamp}.docx",
                mapping=living_replacements,
                output_list=generated_files
            ):
                with open(debug_log_path, 'a') as log:
                    log.write("Generated Living Will doc\n")
            else:
                with open(debug_log_path, 'a') as log:
                    log.write("Failed to generate Living Will doc\n")

        if not generated_files:
            return JsonResponse({'error': 'No documents were generated'}, status=400)

        # Send email with attachments
        try:
            email = EmailMessage(
                subject=f"{first_name} {last_name}'s Generated Documents",
                body='Attached are your generated documents.',
                from_email=settings.EMAIL_HOST_USER,
                to=[data['email']]
            )
            for file_path in generated_files:
                with open(file_path, 'rb') as f:
                    email.attach(os.path.basename(file_path), f.read(), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            email.send()
            with open(debug_log_path, 'a') as log:
                log.write("Email sent successfully\n")
        except Exception as e:
            with open(debug_log_path, 'a') as log:
                log.write(f"Email sending failed: {str(e)}\n{traceback.format_exc()}\n")
            return JsonResponse({'error': f'Failed to send email: {str(e)}'}, status=500)

        # Log generated files
        with open(debug_log_path, 'a') as log:
            log.write("Generated files:\n")
            for path in generated_files:
                log.write(f" - {path}\n")

        # Clean up files to avoid clutter
        for file_path in generated_files:
            if os.path.exists(file_path):
                os.remove(file_path)
        if os.path.exists(user_output_dir):
            try:
                os.rmdir(user_output_dir)
            except OSError:
                pass

        # Return download links
        response = ''.join([f"<p><a href='{f}' download>Download {os.path.basename(f)}</a></p>" for f in generated_files])
        return JsonResponse({'success': True, 'html': response})

    except Exception as e:
        with open(debug_log_path, 'a') as log:
            log.write(f"Error: {str(e)}\n{traceback.format_exc()}\n")
        return JsonResponse({'error': f'Internal server error: {str(e)}'}, status=500)

def index(request):
    return render(request, 'index.html')
